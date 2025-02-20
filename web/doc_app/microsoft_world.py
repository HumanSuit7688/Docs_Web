import docx
import datetime
from docx.shared import Pt
from pytrovich.detector import PetrovichGenderDetector
from pytrovich.enums import NamePart, Gender, Case
from pytrovich.maker import PetrovichDeclinationMaker


def declension(surname, name, patronymic):
    maker = PetrovichDeclinationMaker()
    detector = PetrovichGenderDetector()

    gender = detector.detect(firstname=name)
    if str(gender) == 'Gender.MALE':
        print('парень')
        name_dec = maker.make(NamePart.FIRSTNAME, Gender.MALE, Case.DATIVE, name)
        surname_dec = maker.make(NamePart.LASTNAME, Gender.MALE, Case.DATIVE, surname)
        patronymic_dec = maker.make(NamePart.MIDDLENAME, Gender.MALE, Case.DATIVE, patronymic)
        dict = {'surname': surname_dec, 'name': name_dec, 'patronymic': patronymic_dec}
        return dict
    elif str(gender) == 'Gender.FEMALE':
        print('девушка')
        name_dec = maker.make(NamePart.FIRSTNAME, Gender.FEMALE, Case.DATIVE, name)
        surname_dec = maker.make(NamePart.LASTNAME, Gender.FEMALE, Case.DATIVE, surname)
        patronymic_dec = maker.make(NamePart.MIDDLENAME, Gender.FEMALE, Case.DATIVE, patronymic)
        dict = {'surname': surname_dec, 'name': name_dec, 'patronymic': patronymic_dec}
        return dict


def date_in_par2():
    date_now = datetime.date.today()
    day = date_now.strftime('%d')
    month = date_now.strftime('%m')
    year = date_now.strftime('%Y')

    month_name = {'01' : 'января', '02' : 'февраля', '03' : 'марта', '04' : 'апреля', '05' : 'мая', '06' : 'июня', '07' : 'июля', '08' : 'августа', '09' : 'сентября',
                  '10' : 'октября', '11' : 'ноября', '12' : 'декабря', }
    month = month_name[month]
    text = f'«{day}» {month} {year}г.'
    return text


def fio_in_par4(par4, surname, name, patronymic):
    fio = f'{surname} {name} {patronymic}'
    text = par4.replace('___________________________________________', fio)
    return text


def grade_in_par4(par4, grade_c, grade_b):
    grade = f'{grade_c}{grade_b}'
    text = par4.replace('___', grade)
    return text


def make_doc(surname, name, patronymic, grade_c, grade_b):
    fio_dict = declension(surname, name, patronymic)
    name_name = name
    surname_name = surname
    name = fio_dict.get('name')
    surname = fio_dict.get('surname')
    patronymic = fio_dict.get('patronymic')

    doc = docx.Document('web\doc_app\Pattern.docx')
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    par4 = doc.paragraphs[3].text
    first_step = fio_in_par4(par4, surname, name, patronymic)
    par4_new = grade_in_par4(first_step, grade_c, grade_b)
    doc.paragraphs[1].runs[0].text = date_in_par2()
    doc.paragraphs[3].text = par4_new

    date_now = datetime.date.today()
    day = date_now.strftime('%d')
    month = date_now.strftime('%m')

    doc.save(f'new_docs/new_{surname_name}_{name_name}_{day}_{month}.docx')

# make_doc('Урбан', 'Иван', 'Борисович', 10, 'А')
# print(declension('Урбан', 'Иван', 'Борисович'))

